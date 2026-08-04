"""Build the English factor-level analysis report as an editable Word document.

Reads data/comparisons/l4_factor_analysis/ (CSVs + summary.json + figures_en/)
and writes factor_book_analysis.docx next to them.
"""
from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "data/comparisons/l4_factor_analysis"
FIG = OUT / "figures_en"
DEST = OUT / "factor_book_analysis.docx"

pf = pd.read_csv(OUT / "per_factor_ic.csv")
comb = pd.read_csv(OUT / "combined_book_ic.csv").set_index("combo")
near = pd.read_csv(OUT / "nearest_zoo_corr.csv")
S = json.loads((OUT / "summary.json").read_text())

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

GREY = RGBColor(0x60, 0x5D, 0x54)


def para(text, italic=False, size=None, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullets(items):
    for lead, rest in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(lead)
        r.bold = True
        p.add_run(rest)


def table(headers, rows, widths=None, bold_cols=()):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(9)
            if j in bold_cols:
                r.bold = True
            if j > 0 and isinstance(v, str) and any(ch.isdigit() for ch in v):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def figure(name, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / name), width=Inches(width))
    c = doc.add_paragraph()
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def fmt(v, d=3):
    if v is None or (isinstance(v, float) and v != v):
        return "–"
    return f"{v:.{d}f}"


# ── title ────────────────────────────────────────────────────────────────────
doc.add_heading("Factor-Level Analysis of Two LLM-Evolved Alpha Books", level=0)
para("GPT-5.6 Terra (L4 configuration) vs. Claude Opus 5, benchmarked against the "
     "101 formulaic alphas — Nasdaq-100 point-in-time panel, 2010–2026.",
     italic=True, color=GREY)

# ── 0. what is measured ─────────────────────────────────────────────────────
doc.add_heading("1. What is measured", level=1)
para("Each evolution run publishes a book of factors: small Python programs that map the "
     "market panel (prices, volumes, fundamentals, estimates) to one signal value per "
     "(day, stock). This analysis evaluates those signals directly — no strategy "
     "construction, no position sizing, no costs.")
para("Two kinds of information coefficient (IC) are reported, both against 6-day forward "
     "returns (h = 6 bars, the horizon both runs searched at):")
bullets([
    ("Per-factor IC. ", "For each factor and each stock, the Pearson correlation between "
     "the factor's signal series and that stock's own 6-day forward-return series; the "
     "observation-weighted average across stocks is the factor's pooled per-underlying IC. "
     "This is the statistic family the evolutionary harness itself optimises, so it is the "
     "honest yardstick for the search."),
    ("Combined-book IC. ", "All signals of a book are standardised per stock and fed to a "
     "LightGBM regressor (the same model class the runs used to score marginal value), "
     "fitted once on the DEV window only, predicting the 6-day forward return. The pooled "
     "per-underlying IC of that single combined prediction measures the book as a whole, "
     "including nonlinear interactions between factors."),
])
para("Three evaluation windows, following the runs' own 60/20/20 split convention of the "
     "research panel (2010 → 2024-07):")
table(
    ["Window", "Period", "Bars", "Role"],
    [["DEV (IS+VAL)", f"{S['windows']['dev'][0]} → {S['windows']['dev'][1]}",
      str(S['windows']['dev'][2]),
      "Search window: the evolution selected factors on (parts of) this data; the "
      "combined model is fitted here. Combined ICs on DEV are in-sample."],
     ["TEST", f"{S['windows']['test'][0]} → {S['windows']['test'][1]}",
      str(S['windows']['test'][2]),
      "Final 20% of the research panel, never revealed to any generation of the search."],
     ["FORWARD", f"{S['windows']['forward'][0]} → {S['windows']['forward'][1]}",
      str(S['windows']['forward'][2]),
      "Entirely outside the research panel (data downloaded after the runs finished). "
      "The cleanest out-of-sample evidence."]],
)

p = doc.add_paragraph()
r = p.add_run(
    "Comparability caveat: no Opus 5 run exists at the L4 configuration (the Opus ladder "
    "stopped after L2 when its budget was exhausted). The comparison is therefore "
    "L4_terra_s0 (GPT-5.6, GraphRAG retrieval, 8 mechanism groups, 44 published factors, "
    "$87) vs. L2_opus5_s0 (Claude Opus 5, no retrieval, 1 group × 4 demes, 18 factors, "
    "$452). Model and configuration are confounded; differences cannot be attributed to "
    "the model alone.")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY
para("Four of the 44 Terra factors produce degenerate (variance-free) signals on the "
     "extended analysis panel and are excluded from per-factor and correlation "
     "statistics; they remain in the combined fits, where they contribute nothing.",
     size=9.5, color=GREY, italic=True)

# ── 2. combined books ───────────────────────────────────────────────────────
doc.add_heading("2. Combined books: IC by window", level=1)
order = ["zoo_101", "terra_book", "opus_book", "terra_plus_opus",
         "terra_plus_zoo", "opus_plus_zoo", "terra_plus_opus_zoo"]
label = {"zoo_101": "101 formulaic alphas alone", "terra_book": "Terra book (44)",
         "opus_book": "Opus book (18)", "terra_plus_opus": "Terra + Opus (62)",
         "terra_plus_zoo": "Terra + 101 (145)", "opus_plus_zoo": "Opus + 101 (119)",
         "terra_plus_opus_zoo": "Terra + Opus + 101 (163)"}
rows = []
for c in order:
    if c not in comb.index:
        continue
    r_ = comb.loc[c]
    rows.append([label[c], fmt(r_["ic_pooled_dev"]), fmt(r_["ic_pooled_test"]),
                 fmt(r_["ic_pooled_forward"])])
table(["Book", "DEV (in-sample)", "TEST", "FORWARD"], rows, bold_cols=(2, 3))
figure("fig3_combined_ic.png",
       "Figure 3 — Combined-book IC per window (LightGBM fitted on DEV only). DEV bars are "
       "in-sample fits and are shown only to gauge the generalisation gap.")

t_alone = comb.loc["terra_book"]
o_alone = comb.loc["opus_book"]
z = comb.loc["zoo_101"]
to = comb.loc["terra_plus_opus"] if "terra_plus_opus" in comb.index else None
toz = comb.loc["terra_plus_opus_zoo"] if "terra_plus_opus_zoo" in comb.index else None
obs = [
    ("Opus generalises best on its own. ",
     f"TEST IC {fmt(o_alone['ic_pooled_test'])} from a DEV fit of "
     f"{fmt(o_alone['ic_pooled_dev'])} — a retention of "
     f"{o_alone['ic_pooled_test']/o_alone['ic_pooled_dev']:.0%}, vs. "
     f"{t_alone['ic_pooled_test']/t_alone['ic_pooled_dev']:.0%} for Terra and "
     f"{z['ic_pooled_test']/z['ic_pooled_dev']:.0%} for the formulaic library."),
    ("Terra is a marginal-value book. ",
     f"Alone it sits at library level (TEST {fmt(t_alone['ic_pooled_test'])}), but added "
     f"to the 101 alphas it lifts the FORWARD IC from {fmt(z['ic_pooled_forward'])} to "
     f"{fmt(comb.loc['terra_plus_zoo']['ic_pooled_forward'])} — consistent with the L4 "
     "run's selection objective (leave-one-out marginal contribution against exactly this "
     "reference book)."),
]
if to is not None:
    obs.append(("The two books combine well. ",
                f"Terra + Opus reaches TEST {fmt(to['ic_pooled_test'])} / FORWARD "
                f"{fmt(to['ic_pooled_forward'])} from only 62 factors"
                + (f"; adding the 101 alphas on top gives TEST "
                   f"{fmt(toz['ic_pooled_test'])} / FORWARD {fmt(toz['ic_pooled_forward'])}"
                   if toz is not None else "") + "."))
bullets(obs)

# ── 3. per-factor ICs ───────────────────────────────────────────────────────
doc.add_heading("3. Per-factor ICs: distribution and degradation", level=1)
rows = []
for bk, lb in [("terra", "Terra (L4)"), ("opus", "Opus (L2-evolution)"),
               ("zoo", "101 alphas")]:
    d = pf[pf.book == bk].dropna(subset=["ic_pooled_dev"])
    rows.append([lb, str(len(d)),
                 fmt(d.ic_pooled_dev.abs().mean(), 4), fmt(d.ic_pooled_test.abs().mean(), 4),
                 fmt(d.ic_pooled_forward.abs().mean(), 4),
                 f"{(np.sign(d.ic_pooled_dev) == np.sign(d.ic_pooled_test)).mean():.0%}",
                 f"{d.ic_pooled_dev.corr(d.ic_pooled_test):.2f}"])
table(["Book", "n", "mean |IC| DEV", "mean |IC| TEST", "mean |IC| FWD",
       "sign retention DEV→TEST", "corr(DEV, TEST)"], rows)
figure("fig1_ic_scatter.png",
       "Figure 1 — Each point is one factor: DEV IC against TEST IC (left) and FORWARD IC "
       "(right). Points do not collapse below the diagonal: average per-factor ICs do not "
       "degrade out of sample, because the search never selected on standalone IC (it "
       "selected on marginal contribution to the book), so standalone IC was never "
       "over-optimised. Sign retention is 80–89% for the LLM books vs. 74–75% for the "
       "formulaic library.")

p = doc.add_paragraph()
r = p.add_run("Key finding: the apparent overfitting sits in the combining model, not in "
              "the factors. ")
r.bold = True
p.add_run("The combined-book IC drops by roughly an order of magnitude from DEV to TEST "
          "(e.g. Terra 0.213 → 0.024), which at first glance looks like heavy factor "
          "overfitting. The per-factor evidence contradicts that reading: individual "
          "factor ICs are as strong on TEST and FORWARD as on DEV (mean |IC| flat to "
          "rising, sign retention 80–89%). What collapses out of sample is the fitted "
          "LightGBM combination — its in-sample IC reflects the model memorising DEV-"
          "specific interaction patterns, an in-sample artefact that says little about "
          "the factors themselves. The honest generalisation gap of the books is the "
          "much smaller one visible per factor, plus the 2021 regime break documented "
          "in Section 7.")

# ── 4. best factors ─────────────────────────────────────────────────────────
doc.add_heading("4. Strongest individual factors", level=1)


def top_rows(bk):
    d = pf[pf.book == bk].dropna(subset=["ic_pooled_dev"])
    d = d.reindex(d.ic_pooled_test.abs().sort_values(ascending=False).index).head(8)
    return [[r.factor_id, (r.category if isinstance(r.category, str) else "–"),
             fmt(r.ic_pooled_dev, 4), fmt(r.ic_pooled_test, 4),
             fmt(r.ic_pooled_forward, 4)] for r in d.itertuples()]


para("Opus — top 8 by |TEST IC|:", bold=True)
table(["Factor", "Category", "IC DEV", "IC TEST", "IC FWD"], top_rows("opus"),
      bold_cols=(3,))
para("Terra — top 8 by |TEST IC|:", bold=True)
table(["Factor", "Category", "IC DEV", "IC TEST", "IC FWD"], top_rows("terra"),
      bold_cols=(3,))
dz = pf[pf.book == "zoo"]
best_zoo = dz.reindex(dz.ic_pooled_test.abs().sort_values(ascending=False).index).head(2)
para(f"Reference: the best formulaic alphas on TEST are "
     f"{best_zoo.iloc[0].factor_id} ({fmt(best_zoo.iloc[0].ic_pooled_test, 3)}) and "
     f"{best_zoo.iloc[1].factor_id} ({fmt(best_zoo.iloc[1].ic_pooled_test, 3)}). The single "
     "best factor of the whole analysis is Opus's comparable_basket_repair_rate_activity "
     "(TEST 0.062 from a DEV of only 0.024).")
figure("fig2_top_factors_opus.png", "Figure 2a — Opus: top 15 factors by |DEV IC|, with "
       "TEST and FORWARD alongside.")
figure("fig2_top_factors_terra.png", "Figure 2b — Terra: top 15 factors by |DEV IC|.")

# ── 5. correlation & diversity ──────────────────────────────────────────────
doc.add_heading("5. Correlation structure and diversity", level=1)
para(f"Signal correlations are computed on cross-sectionally standardised signals over the "
     f"DEV window. Within-book redundancy is low for both LLM books (mean |ρ|: Terra "
     f"{S['terra_mean_abs_corr']:.3f}, Opus {S['opus_mean_abs_corr']:.3f}, library "
     f"{S['zoo_mean_abs_corr']:.3f}), but efficiency differs: the participation ratio of "
     f"the correlation spectrum gives Opus {S['opus_eff_n']:.1f} effective factors out of "
     f"{S['opus_n']} (71%), Terra {S['terra_eff_n']:.1f} / {S['terra_n']} (48%), and the "
     f"formulaic library only {S['zoo_eff_n']:.1f} / {S['zoo_n']} (26%).")
para(f"The cross-blocks are weak: mean |ρ| Terra×Opus {S['terra_x_opus_mean_abs_corr']:.3f} "
     f"(max {S['terra_x_opus_max_abs_corr']:.2f}), Terra×library "
     f"{S['terra_x_zoo_mean_abs_corr']:.3f}, Opus×library "
     f"{S['opus_x_zoo_mean_abs_corr']:.3f}. The two models explored largely orthogonal "
     f"factor spaces: Opus concentrated on microstructure (12 of 18 factors), Terra spread "
     f"across sentiment, volatility, fundamentals and microstructure. Pooling everything "
     f"yields {S['eff_n_terra+opus+zoo']:.1f} effective factors from 160.")
figure("fig4_corr_heatmaps.png",
       "Figure 4 — Within-book correlations, hierarchically ordered. Terra carries one "
       "visible redundancy cluster (liquidity/downside family, pairs up to |ρ| ≈ 0.85); "
       "Opus a small four-factor wick/absorption block.")
figure("fig5_corr_blocks.png",
       "Figure 5 — Full correlation matrix with block boundaries. The library block "
       "(bottom right) is visibly denser; the off-diagonal blocks stay pale.")
figure("fig7_effective_n.png",
       "Figure 7 — Nominal vs. effective factor counts (participation ratio of the "
       "correlation eigenvalues).")

# ── 6. novelty ──────────────────────────────────────────────────────────────
doc.add_heading("6. Novelty relative to the formulaic library", level=1)
med_t = near[near.book == "terra"].max_abs_corr_zoo.median()
med_o = near[near.book == "opus"].max_abs_corr_zoo.median()
para(f"For each evolved factor, the maximum |correlation| to any of the 101 formulaic "
     f"alphas measures how much of it the library already spans. The median is "
     f"{med_t:.2f} (Terra) and {med_o:.2f} (Opus): the typical evolved factor is not a "
     f"disguised formulaic alpha. The most library-like factors of both books attach to "
     f"the same two alphas (alpha_033, alpha_038), with a maximum of 0.78.")
nr = near.sort_values("max_abs_corr_zoo", ascending=False).head(6)
table(["Book", "Factor", "Nearest alpha", "max |ρ| library", "max |ρ| own book"],
      [[("Terra" if r.book == "terra" else "Opus"), r.factor_id, r.nearest_alpha,
        fmt(r.max_abs_corr_zoo, 2), fmt(r.max_abs_corr_own_book, 2)]
       for r in nr.itertuples()])
figure("fig6_zoo_novelty.png",
       "Figure 6 — Left: distribution of library-similarity (binned in 0.1-wide buckets); "
       "Terra has a solid block of genuinely novel factors below 0.15. Right: factors above "
       "the diagonal are more redundant within their own book than with the library — "
       "Terra's redundancy is home-grown (the cluster of Figure 4), not library imitation.")

# ── 7. IC through time ──────────────────────────────────────────────────────
doc.add_heading("7. IC through time: the 2021 regime break", level=1)
para("The pooled per-underlying IC is one number per window and has no daily analogue, so "
     "the time-resolved view uses the daily cross-sectional IC (each day, the correlation "
     "across stocks between signal and forward return) as a diagnostic; its level is not "
     "comparable to the pooled numbers above.")
figure("fig8_rolling_ic.png",
       "Figure 8 — 63-day rolling mean of the daily cross-sectional IC of the combined "
       "signals (fitted on DEV only). All three books — including the formulaic library, "
       "which was never fitted to this panel — collapse simultaneously at the 2021 TEST "
       "boundary: the cross-sectional factor regime in the Nasdaq-100 itself changed "
       "(mega-cap concentration), so the DEV→TEST drop is not pure overfitting. The Opus "
       "signal stays above zero most often out of sample; the library turns clearly "
       "negative in 2025–26.")

# ── 8. conclusions ──────────────────────────────────────────────────────────
doc.add_heading("8. Conclusions", level=1)
bullets([
    ("Per-factor quality: Opus. ", "Individually stronger and more stable factors "
     "(mean |TEST IC| 0.018 vs. 0.011), the best standalone generalisation of the combined "
     "book (TEST 0.045), and the highest effective-dimension efficiency — at 5× the cost "
     "($452 vs. $87) and with only 18 factors."),
    ("Breadth and marginal value: Terra. ", "More factors, more genuine novelty, the "
     "largest diversity gain to the pooled book, and the best FORWARD results appear only "
     "in combination (Terra + 101: 0.041) — exactly the selection objective of the L4 "
     "configuration."),
    ("Near-orthogonal books. ", "Mean cross-correlation of 0.056 between the two books; "
     "combining them (with or without the formulaic library) is the natural next step and "
     "already outperforms either book alone on TEST."),
    ("The DEV→TEST gap is a model artefact, not factor overfitting. ", "Per-factor ICs "
     "hold or improve out of sample while only the fitted combination collapses; "
     "in-sample combined ICs should never be quoted as evidence of book quality."),
    ("Attribution caveat. ", "Model (GPT-5.6 vs. Opus 5) and configuration (L4 GraphRAG / "
     "8 groups vs. L2 no-retrieval / 1 group) are confounded; a clean model comparison "
     "would require the never-run L4 Opus arm."),
])

para("Reproduction: scripts/analyze_l4_factor_books.py → data/comparisons/"
     "l4_factor_analysis/ (CSVs, summary.json), figures: scripts/plot_l4_factor_books_en.py.",
     size=9, color=GREY)

doc.save(DEST)
print("wrote", DEST)
